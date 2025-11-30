# -*- coding: utf-8 -*-
"""
Виджет поиска AI для ApiAi
"""

import os
import json
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QPushButton, 
    QLineEdit, QTextBrowser, QLabel, QComboBox, 
    QMessageBox, QProgressBar, QSplitter, QFrame,
    QFileDialog
)
from PySide6.QtCore import Qt, QThread, Signal, QTimer
from PySide6.QtGui import QFont, QIcon

from .pdf_search import AIPDFSearcher


class AISearchWorker(QThread):
    """Worker для AI поиска в отдельном потоке"""
    finished = Signal(dict)
    
    def __init__(self, provider: str, api_key: str, query: str, api_url: str = None, 
                 custom_prompt: str = None, use_encryption: bool = False, encryption_key: str = None,
                 app_id: str = "bomcategorizer-v5"):
        super().__init__()
        self.provider = provider
        self.api_key = api_key
        self.query = query
        self.api_url = api_url
        self.custom_prompt = custom_prompt
        self.use_encryption = use_encryption
        self.encryption_key = encryption_key
        self.app_id = app_id
    
    def run(self):
        """Выполняет AI поиск"""
        searcher = AIPDFSearcher(
            self.provider, 
            self.api_key, 
            self.api_url,
            use_encryption=self.use_encryption,
            encryption_key=self.encryption_key,
            app_id=self.app_id
        )
        
        # Используем кастомный промпт если передан
        if self.custom_prompt:
            results = searcher.search_with_prompt(self.query, self.custom_prompt)
        else:
            results = searcher.search(self.query)
        
        self.finished.emit(results)


class SearchWidget(QWidget):
    """Виджет для взаимодействия с AI API"""
    
    def __init__(self, parent=None, config=None):
        super().__init__(parent)
        self.config = config or {}
        self.parent_window = parent
        
        self._create_ui()
        
    def _create_ui(self):
        """Создает интерфейс виджета"""
        layout = QVBoxLayout(self)
        layout.setContentsMargins(10, 10, 10, 10)
        layout.setSpacing(10)
        
        # --- Верхняя панель: Выбор провайдера и настройки ---
        top_panel = QHBoxLayout()
        
        self.provider_combo = QComboBox()
        self.provider_combo.addItems(["Anthropic Claude", "OpenAI GPT", "Telegram Bot"])
        self.provider_combo.setToolTip("Выберите провайдера AI")
        self.provider_combo.currentIndexChanged.connect(self._on_provider_changed)
        top_panel.addWidget(QLabel("Провайдер:"))
        top_panel.addWidget(self.provider_combo)
        
        top_panel.addStretch()
        
        # Кнопка настроек (если нужно, но у нас есть меню)
        # settings_btn = QPushButton("⚙️ Настройки")
        # settings_btn.clicked.connect(self.open_settings)
        # top_panel.addWidget(settings_btn)
        
        layout.addLayout(top_panel)
        
        # --- Панель поиска ---
        search_panel = QHBoxLayout()
        
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Введите название компонента или запрос...")
        self.search_input.returnPressed.connect(self.start_search)
        search_panel.addWidget(self.search_input)
        
        self.search_btn = QPushButton("🔍 Поиск")
        self.search_btn.clicked.connect(self.start_search)
        self.search_btn.setStyleSheet("""
            QPushButton {
                background-color: #007AFF;
                color: white;
                font-weight: bold;
                padding: 5px 15px;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #0062CC;
            }
        """)
        search_panel.addWidget(self.search_btn)
        
        layout.addLayout(search_panel)
        
        # --- Область промпта (складная или всегда видимая) ---
        prompt_group = QFrame()
        prompt_group.setFrameShape(QFrame.StyledPanel)
        prompt_layout = QVBoxLayout(prompt_group)
        prompt_layout.setContentsMargins(5, 5, 5, 5)
        
        prompt_header = QHBoxLayout()
        prompt_header.addWidget(QLabel("📝 Промпт:"))
        
        self.prompt_type_combo = QComboBox()
        self.prompt_type_combo.addItems(["Стандартный (информация о компоненте)", "Пользовательский"])
        self.prompt_type_combo.currentIndexChanged.connect(self._on_prompt_type_changed)
        prompt_header.addWidget(self.prompt_type_combo)
        
        prompt_header.addStretch()
        prompt_layout.addLayout(prompt_header)
        
        self.custom_prompt_edit = QTextBrowser() # Используем QTextBrowser для отображения, если стандартный
        self.custom_prompt_edit.setReadOnly(True)
        self.custom_prompt_edit.setMaximumHeight(100)
        self.custom_prompt_edit.setPlaceholderText("Здесь будет отображаться используемый промпт...")
        prompt_layout.addWidget(self.custom_prompt_edit)
        
        # Делаем редактируемым только для пользовательского режима
        self.custom_prompt_edit.setReadOnly(True)
        
        layout.addWidget(prompt_group)
        
        # --- Область результатов ---
        self.results_browser = QTextBrowser()
        self.results_browser.setOpenExternalLinks(True)
        layout.addWidget(self.results_browser)
        
        # --- Прогресс бар ---
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)
        
        # --- Нижняя панель действий ---
        actions_panel = QHBoxLayout()
        
        save_btn = QPushButton("💾 Сохранить в файл")
        save_btn.clicked.connect(self.save_results)
        actions_panel.addWidget(save_btn)
        
        actions_panel.addStretch()
        
        clear_btn = QPushButton("🧹 Очистить")
        clear_btn.clicked.connect(self.clear_results)
        actions_panel.addWidget(clear_btn)
        
        layout.addLayout(actions_panel)
        
        # Инициализация состояния
        self._update_prompt_preview()

    def _on_provider_changed(self):
        """Обработчик смены провайдера"""
        # Можно добавить логику проверки наличия ключа
        pass

    def _on_prompt_type_changed(self):
        """Обработчик смены типа промпта"""
        is_custom = self.prompt_type_combo.currentIndex() == 1
        
        if is_custom:
            self.custom_prompt_edit.setReadOnly(False)
            self.custom_prompt_edit.setText("")
            self.custom_prompt_edit.setPlaceholderText("Введите ваш промпт здесь. Используйте {query} для подстановки запроса.")
            # Меняем тип на QTextEdit для редактирования (хак, лучше пересоздать или иметь два виджета)
            # Но для простоты пока оставим QTextBrowser и сделаем setReadOnly(False) - это работает для QTextEdit, 
            # но QTextBrowser наследует от QTextEdit и обычно read-only.
            # Проверим: QTextBrowser обычно read-only по дизайну.
            # Лучше заменить на QTextEdit.
            
            # Замена виджета на лету сложна, поэтому лучше изначально иметь QTextEdit
            pass 
        else:
            self.custom_prompt_edit.setReadOnly(True)
            self._update_prompt_preview()
            
    def _update_prompt_preview(self):
        """Обновляет предпросмотр стандартного промпта"""
        if self.prompt_type_combo.currentIndex() == 0:
            self.custom_prompt_edit.setText(
                "Найди информацию об электронном компоненте: {query}\n\n"
                "Пожалуйста, предоставь следующую информацию в структурированном виде:\n"
                "1. Полное название и производитель\n"
                "2. Тип компонента\n"
                "3. Основные характеристики\n"
                "4. Краткое описание назначения\n"
                "5. Типичные примеры использования\n"
                "6. Прямая ссылка на PDF документацию"
            )

    def start_search(self):
        """Запускает поиск"""
        query = self.search_input.text().strip()
        if not query:
            QMessageBox.warning(self, "Ошибка", "Введите поисковый запрос")
            return
            
        # Проверка ключей
        api_keys = self.config.get("api_keys", {})
        provider_idx = self.provider_combo.currentIndex()
        
        api_key = ""
        provider_code = ""
        api_url = None
        use_encryption = False
        encryption_key = None
        
        if provider_idx == 0: # Anthropic
            api_key = api_keys.get("anthropic", "")
            provider_code = "anthropic"
            if not api_key:
                QMessageBox.warning(self, "Ошибка", "Не задан API ключ для Anthropic Claude")
                return
        elif provider_idx == 1: # OpenAI
            api_key = api_keys.get("openai", "")
            provider_code = "openai"
            if not api_key:
                QMessageBox.warning(self, "Ошибка", "Не задан API ключ для OpenAI GPT")
                return
        elif provider_idx == 2: # Telegram Bot
            api_key = api_keys.get("telegram_key", "")
            provider_code = "telegram_bot"
            api_url = api_keys.get("telegram_url", "")
            use_encryption = api_keys.get("telegram_use_encryption", True)
            encryption_key = api_keys.get("telegram_enc_key", "")
            
            if not api_url:
                QMessageBox.warning(self, "Ошибка", "Не задан URL для Telegram Bot API")
                return
        
        # Подготовка промпта
        custom_prompt = None
        if self.prompt_type_combo.currentIndex() == 1: # Custom
            # Для кастомного промпта мы берем текст из поля, но нам нужно подставить query
            # Однако, AISearchWorker ожидает, что если custom_prompt передан, то он уже готовый?
            # Или он форматирует? В AIPDFSearcher.search_with_prompt просто отправляет prompt.
            # Значит нам нужно сформировать полный промпт здесь.
            
            raw_prompt = self.custom_prompt_edit.toPlainText()
            if "{query}" in raw_prompt:
                custom_prompt = raw_prompt.replace("{query}", query)
            else:
                # Если нет плейсхолдера, просто добавляем запрос в начало или конец?
                # Лучше считать, что пользователь сам написал всё что нужно.
                # Но если он просто написал "расскажи про...", то query уже внутри?
                # Давайте просто передадим как есть, если пользователь ввел текст.
                # Но если текст пустой, то ошибка.
                if not raw_prompt.strip():
                    QMessageBox.warning(self, "Ошибка", "Введите текст пользовательского промпта")
                    return
                custom_prompt = raw_prompt
        
        # UI update
        self.search_btn.setEnabled(False)
        self.progress_bar.setVisible(True)
        self.progress_bar.setRange(0, 0) # Indeterminate
        self.results_browser.clear()
        self.results_browser.append(f"⏳ Выполняется поиск: {query}...")
        
        # Запуск воркера
        self.worker = AISearchWorker(
            provider=provider_code,
            api_key=api_key,
            query=query,
            api_url=api_url,
            custom_prompt=custom_prompt,
            use_encryption=use_encryption,
            encryption_key=encryption_key
        )
        self.worker.finished.connect(self._on_search_finished)
        self.worker.start()
        
    def _on_search_finished(self, results):
        """Обработка результатов поиска"""
        self.search_btn.setEnabled(True)
        self.progress_bar.setVisible(False)
        
        if not results:
            self.results_browser.setHtml("<h3 style='color: red'>Ошибка: Пустой ответ</h3>")
            return
            
        if results.get('error'):
            self.results_browser.setHtml(f"<h3 style='color: red'>Ошибка: {results['error']}</h3>")
            if 'raw_response' in results:
                self.results_browser.append("\n--- Raw Response ---\n")
                self.results_browser.append(results['raw_response'])
            return
            
        # Форматирование вывода
        html = self._format_results_html(results)
        self.results_browser.setHtml(html)
        
    def _format_results_html(self, data):
        """Форматирует результаты в HTML"""
        component = data.get('component', 'Неизвестный компонент')
        provider = data.get('provider', 'AI')
        
        html = f"""
        <h2 style='color: #2c3e50'>Результаты поиска: {component}</h2>
        <div style='color: #7f8c8d; font-size: 12px'>Провайдер: {provider}</div>
        <hr>
        """
        
        if data.get('found'):
            html += f"""
            <p><b>Полное название:</b> {data.get('full_name', '-')}</p>
            <p><b>Производитель:</b> {data.get('manufacturer', '-')}</p>
            <p><b>Тип:</b> {data.get('type', '-')}</p>
            
            <h3>Описание</h3>
            <p>{data.get('description', '-')}</p>
            
            <h3>Характеристики</h3>
            <ul>
            """
            
            specs = data.get('specifications', {})
            if isinstance(specs, dict):
                for k, v in specs.items():
                    html += f"<li><b>{k}:</b> {v}</li>"
            
            html += "</ul><h3>Примеры использования</h3><ul>"
            
            examples = data.get('examples', [])
            if isinstance(examples, list):
                for ex in examples:
                    html += f"<li>{ex}</li>"
            
            html += "</ul>"
            
            url = data.get('datasheet_url')
            if url:
                html += f"<p><b>Datasheet:</b> <a href='{url}'>{url}</a></p>"
                
        else:
            # Если просто текст (например, для кастомного промпта без JSON)
            if 'description' in data:
                html += f"<pre>{data['description']}</pre>"
            elif 'raw_response' in data:
                html += f"<pre>{data['raw_response']}</pre>"
            else:
                html += "<p>Информация не найдена.</p>"
                
        return html

    def clear_results(self):
        """Очищает результаты"""
        self.results_browser.clear()
        self.search_input.clear()
        self.search_input.setFocus()

    def save_results(self):
        """Сохраняет результаты в файл"""
        content = self.results_browser.toPlainText()
        if not content:
            QMessageBox.warning(self, "Пусто", "Нет результатов для сохранения")
            return
            
        filters = (
            "HTML Files (*.html);;"
            "Text Files (*.txt);;"
            "Word Document (*.docx);;"
            "PDF (*.pdf)"
        )
        
        file_path, selected_filter = QFileDialog.getSaveFileName(
            self, "Сохранить результаты", "", filters
        )
        
        if not file_path:
            return
            
        try:
            if file_path.endswith('.html'):
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(self.results_browser.toHtml())
                    
            elif file_path.endswith('.txt'):
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                    
            elif file_path.endswith('.docx'):
                self._save_as_docx(file_path, content)
                
            elif file_path.endswith('.pdf'):
                self._save_as_pdf(file_path, content)
            else:
                # Default to txt
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                    
            QMessageBox.information(self, "Успех", f"Файл сохранен:\n{file_path}")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить файл:\n{e}")
    
    def _save_as_docx(self, file_path, content):
        """Сохраняет в Word документ"""
        try:
            from docx import Document
            from docx.shared import Pt
            
            doc = Document()
            
            # Заголовок
            heading = doc.add_heading('Результаты поиска ApiAi', 0)
            heading.style.font.name = 'Arial'
            
            # Контент
            for paragraph in content.split('\n\n'):
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph.strip())
                    p.style.font.name = 'Arial'
                    p.style.font.size = Pt(11)
            
            doc.save(file_path)
            
        except ImportError:
            QMessageBox.warning(
                self, 
                "Модуль не найден",
                "Для сохранения в Word установите: pip install python-docx"
            )
    
    def _save_as_pdf(self, file_path, content):
        """Сохраняет в PDF с поддержкой кириллицы"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle
            from reportlab.lib.units import mm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            import os
            
            # Регистрация шрифта DejaVu Sans (с поддержкой кириллицы)
            font_registered = False
            font_name = 'DejaVuSans'
            
            # Пытаемся найти и зарегистрировать шрифт
            font_paths = [
                os.path.join(os.path.dirname(os.path.dirname(__file__)), 'fonts', 'DejaVuSans.ttf'),
                '/System/Library/Fonts/Supplemental/Arial.ttf',  # macOS
                'C:\\Windows\\Fonts\\arial.ttf',  # Windows
                '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',  # Linux
            ]
            
            for font_path in font_paths:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont(font_name, font_path))
                        font_registered = True
                        break
                    except Exception:
                        continue
            
            if not font_registered:
                # Fallback на Helvetica (без кириллицы)
                font_name = 'Helvetica'
                QMessageBox.warning(
                    self,
                    "Шрифт не найден",
                    "Не найден шрифт с поддержкой кириллицы. PDF может отображаться некорректно."
                )
            
            # Создание документа
            doc = SimpleDocTemplate(
                file_path,
                pagesize=A4,
                leftMargin=15*mm,
                rightMargin=15*mm,
                topMargin=15*mm,
                bottomMargin=15*mm
            )
            
            # Стили
            title_style = ParagraphStyle(
                'Title',
                fontName=font_name,
                fontSize=14,
                leading=16,
                spaceAfter=10
            )
            
            body_style = ParagraphStyle(
                'Body',
                fontName=font_name,
                fontSize=10,
                leading=12,
                spaceAfter=6
            )
            
            # Сборка контента
            story = []
            
            # Заголовок
            story.append(Paragraph("Результаты поиска ApiAi", title_style))
            story.append(Spacer(1, 10*mm))
            
            # Текст результатов
            for paragraph in content.split('\n\n'):
                if paragraph.strip():
                    # Экранируем HTML-спецсимволы
                    safe_text = paragraph.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(safe_text, body_style))
                    story.append(Spacer(1, 3*mm))
            
            # Генерация PDF
            doc.build(story)
            
        except ImportError:
            QMessageBox.warning(
                self,
                "Модуль не найден",
                "Для сохранения в PDF установите: pip install reportlab"
            )

